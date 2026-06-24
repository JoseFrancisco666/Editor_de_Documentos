import os
import streamlit as st
from docx import Document
from pypdf import PdfReader
from striprtf.striprtf import rtf_to_text
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph as ReportLabParagraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# =====================================================================
# SIMULACIÓN DE BASE DE DATOS LOCAL (Temporal hasta conectar Supabase)
# =====================================================================
if "db_usuarios" not in st.session_state:
    # Se deja el diccionario completamente vacío para iniciar desde cero.
    # Los usuarios se crearán en vivo desde la pantalla de Registro.
    st.session_state["db_usuarios"] = {}

# =====================================================================
# INTERFAZ DE GESTIÓN DE ACCESO (LOGIN, REGISTRO, RECUPERACIÓN)
# =====================================================================
def login_privado():
    """Maneja todo el flujo de pantallas para los clientes."""
    if "vista_acceso" not in st.session_state:
        st.session_state["vista_acceso"] = "login"

    # ------------------------- PANTALLA: LOGIN -------------------------
    if st.session_state["vista_acceso"] == "login":
        st.markdown("<h2 style='text-align: center; color: #4A90E2;'>🔐 Acceso al Editor Profesional</h2>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #7F8C8D;'>Ingresa con tu cuenta para comenzar a procesar archivos.</p>", unsafe_allow_html=True)
        
        with st.form("formulario_login"):
            usuario = st.text_input("👤 Usuario o Correo:", placeholder="Escribe tu usuario").strip()
            contrasena = st.text_input("🔑 Contraseña:", type="password", placeholder="Escribe tu contraseña")
            boton_entrar = st.form_submit_button("Ingresar al Sistema")
            
            if boton_entrar:
                usuarios = st.session_state["db_usuarios"]
                if usuario in usuarios and usuarios[usuario] == contrasena:
                    st.session_state["autenticado"] = True
                    st.session_state["usuario_actual"] = usuario
                    st.rerun()
                else:
                    st.error("❌ Usuario o contraseña incorrectos. Verifica tus datos.")

        col_reg, col_olv = st.columns(2)
        with col_reg:
            if st.button("✨ ¿No tienes cuenta? Regístrate aquí"):
                st.session_state["vista_acceso"] = "registro"
                st.rerun()
        with col_olv:
            if st.button("❓ Olvidé mi contraseña"):
                st.session_state["vista_acceso"] = "recuperar"
                st.rerun()

    # ------------------------- PANTALLA: REGISTRO -------------------------
    elif st.session_state["vista_acceso"] == "registro":
        st.markdown("<h2 style='text-align: center; color: #2ECC71;'>📝 Crear una Cuenta Nueva</h2>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #7F8C8D;'>Forma parte de la plataforma y obtén tu espacio privado.</p>", unsafe_allow_html=True)
        
        with st.form("formulario_registro"):
            nuevo_usuario = st.text_input("👤 Elige un Nombre de Usuario:", placeholder="Ejemplo: juan25").strip()
            nueva_contra = st.text_input("🔑 Crea una Contraseña Segura:", type="password", placeholder="Mínimo 6 caracteres")
            confirmar_contra = st.text_input("🔄 Confirma tu Contraseña:", type="password", placeholder="Repite tu contraseña")
            boton_crear = st.form_submit_button("Registrar Cuenta")
            
            if boton_crear:
                if not nuevo_usuario or not nueva_contra:
                    st.warning("⚠️ Todos los campos son obligatorios.")
                elif nueva_contra != confirmar_contra:
                    st.error("❌ Las contraseñas no coinciden. Intenta de nuevo.")
                elif nuevo_usuario in st.session_state["db_usuarios"]:
                    st.error("❌ Este usuario ya se encuentra registrado por otra persona.")
                else:
                    st.session_state["db_usuarios"][nuevo_usuario] = nueva_contra
                    st.success("🎉 ¡Cuenta creada con éxito! Ahora puedes iniciar sesión.")
                    st.session_state["vista_acceso"] = "login"
                    st.rerun()
                    
        if st.button("⬅️ Volver al Inicio de Sesión"):
            st.session_state["vista_acceso"] = "login"
            st.rerun()

    # ------------------------- PANTALLA: RECUPERACIÓN -------------------------
    elif st.session_state["vista_acceso"] == "recuperar":
        st.markdown("<h2 style='text-align: center; color: #E67E22;'>🔄 Recuperar Contraseña</h2>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #7F8C8D;'>Introduce tu usuario para restablecer tus credenciales de acceso.</p>", unsafe_allow_html=True)
        
        with st.form("formulario_recuperar"):
            usuario_olvido = st.text_input("👤 Escribe tu Usuario:", placeholder="El usuario que deseas recuperar").strip()
            boton_buscar = st.form_submit_button("Solicitar Clave")
            
            if boton_buscar:
                if usuario_olvido in st.session_state["db_usuarios"]:
                    st.info(f"🔑 Soporte del sistema: La contraseña de tu cuenta es: **{st.session_state['db_usuarios'][usuario_olvido]}**")
                else:
                    st.error("❌ No encontramos ningún usuario registrado con ese nombre.")
                    
        if st.button("⬅️ Volver al Inicio de Sesión"):
            st.session_state["vista_acceso"] = "login"
            st.rerun()

# =====================================================================
# MOTORES DE LECTURA Y EXTRACCIÓN DE TEXTO PROFUNDA
# =====================================================================
def buscar_cuadros_de_texto(doc):
    from docx.oxml.ns import qn
    textos_escondidos = []
    for elemento in doc.element.iter():
        if elemento.tag.endswith('txbxContent'):
            for p_nodo in elemento.iter(qn('w:p')):
                from docx.text.paragraph import Paragraph
                p = Paragraph(p_nodo, doc)
                if p.text.strip():
                    textos_escondidos.append(p.text.strip())
    return textos_escondidos

def extraer_parrafos_bytes(archivo_subido, nombre_archivo):
    ext = os.path.splitext(nombre_archivo)[1].lower()
    parrafos = []

    if ext in [".docx", ".docm", ".dotx", ".doc"]:
        try:
            doc = Document(archivo_subido)
            for p in doc.paragraphs:
                if p.text.strip():
                    parrafos.append(p.text.strip())
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for p in celda.paragraphs:
                            txt = p.text.strip()
                            if txt and txt not in parrafos:
                                parrafos.append(txt)
            cuadros = buscar_cuadros_de_texto(doc)
            for texto_cuadro in cuadros:
                if texto_cuadro not in parrafos:
                    parrafos.append(texto_cuadro)
        except Exception:
            archivo_subido.seek(0)
            lineas = archivo_subido.read().decode("utf-8", errors="ignore").splitlines()
            parrafos = [l.strip() for l in lineas if l.strip()]

    elif ext == ".pdf":
        reader = PdfReader(archivo_subido)
        for pagina in reader.pages:
            texto_pag = pagina.extract_text()
            if texto_pag:
                bloques = texto_pag.split("\n")
                for b in bloques:
                    if b.strip():
                        parrafos.append(b.strip())

    elif ext == ".rtf":
        contenido_rtf = archivo_subido.read().decode("utf-8", errors="ignore")
        texto_plano = rtf_to_text(contenido_rtf)
        parrafos = [b.strip() for b in texto_plano.split("\n") if b.strip()]

    elif ext in [".txt", ".xml", ".htm", ".html", ".odt"]:
        archivo_subido.seek(0)
        lineas = archivo_subido.read().decode("utf-8", errors="ignore").splitlines()
        parrafos = [l.strip() for l in lineas if l.strip()]

    return parrafos

# =====================================================================
# MOTORES DE COMPILACIÓN DE SALIDA (CONVERSIÓN A BYTES)
# =====================================================================
def exportar_a_docx(parrafos):
    import io
    output = io.BytesIO()
    doc = Document()
    for p in parrafos:
        doc.add_paragraph(p)
    doc.save(output)
    return output.getvalue()

def exportar_a_pdf(parrafos):
    """Genera un archivo PDF sanitizando caracteres de escape XML/HTML."""
    import io
    output = io.BytesIO()
    doc_pdf = SimpleDocTemplate(output, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    estilos = getSampleStyleSheet()
    estilo = ParagraphStyle('Formal', parent=estilos['Normal'], fontName='Helvetica', fontSize=11, leading=16, spaceAfter=12)
    
    historia = []
    for p in parrafos:
        texto_seguro = p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        historia.append(ReportLabParagraph(texto_seguro, estilo))
        
    doc_pdf.build(historia)
    return output.getvalue()

def exportar_a_txt(parrafos):
    contenido = "\n\n".join(parrafos)
    return contenido.encode("utf-8")

def exportar_a_xml(parrafos):
    lineas = ['<?xml version="1.0" encoding="UTF-8"?>', '<documento>']
    for i, p in enumerate(parrafos, 1):
        txt = p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        lineas.append(f'    <parrafo id="{i}">{txt}</parrafo>')
    lineas.append('</documento>')
    return "\n".join(lineas).encode("utf-8")

def exportar_a_html(parrafos):
    lineas = [
        '<!DOCTYPE html>\n<html>\n<head>\n<meta charset="UTF-8">\n<title>Exportado</title>',
        '<style>body{font-family:Arial;margin:40px;line-height:1.6;}p{margin-bottom:15px;}</style>\n</head>\n<body>'
    ]
    for p in parrafos:
        lineas.append(f'    <p>{p}</p>')
    lineas.append('</body>\n</html>')
    return "\n".join(lineas).encode("utf-8")

# =====================================================================
# INTERFAZ DE USUARIO WEB PRINCIPAL (STREAMLIT UI)
# =====================================================================
def main():
    st.set_page_config(page_title="Editor Universal Privado", page_icon="📝", layout="wide")

    if "autenticado" not in st.session_state:
        st.session_state["autenticado"] = False

    if not st.session_state["autenticado"]:
        login_privado()
        return

    # --- PANEL LATERAL DE USUARIO ---
    st.sidebar.markdown("### 👤 Sesión Activa")
    st.sidebar.info(f"Conectado como: **{st.session_state['usuario_actual']}**")
    if st.sidebar.button("🔒 Cerrar Sesión Privada"):
        st.session_state["autenticado"] = False
        st.session_state["usuario_actual"] = None
        st.rerun()

    # --- CUERPO PRINCIPAL DE LA APP ---
    st.markdown("<h1 style='text-align: center; color: #2C3E50;'>📝 Editor e Intercambiador Multiformato</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #7F8C8D;'>Carga un archivo, edita bloques específicos de texto en pantalla y expórtalo al formato de oficina que prefieras.</p>", unsafe_allow_html=True)
    st.divider()

    archivo_subido = st.file_uploader("📥 Arrastra aquí tu archivo (.docx, .pdf, .txt, .xml, .rtf)", type=["docx", "pdf", "txt", "xml", "rtf", "docm", "dotx", "doc", "odt"])

    if archivo_subido is not None:
        if "parrafos_web" not in st.session_state or st.session_state.get("archivo_actual") != archivo_subido.name:
            st.session_state["parrafos_web"] = extraer_parrafos_bytes(archivo_subido, archivo_subido.name)
            st.session_state["archivo_actual"] = archivo_subido.name

        parrafos = st.session_state["parrafos_web"]

        if not parrafos:
            st.error("❌ No se pudo extraer texto de este archivo o se encuentra vacío.")
            return

        col1, col2 = st.columns([1, 1])

        with col1:
            st.subheader("🔍 Estructura Detectada")
            st.caption("Selecciona el párrafo o bloque que deseas modificar:")
            opciones_bloques = [f"[{i}] {p[:70]}..." if len(p) > 70 else f"[{i}] {p}" for i, p in enumerate(parrafos, 1)]
            bloque_seleccionado = st.selectbox("🔢 Elige el bloque a editar:", opciones_bloques)
            idx = int(bloque_seleccionado.split("]")[0].replace("[", "")) - 1

        with col2:
            st.subheader("✍️ Panel de Modificación")
            texto_actual = parrafos[idx]
            nuevo_texto = st.text_area("📄 Contenido del bloque:", value=texto_actual, height=150)
            
            if nuevo_texto != texto_actual:
                parrafos[idx] = nuevo_texto
                st.session_state["parrafos_web"] = parrafos
                st.success("✅ Estructura actualizada en vivo.")

        st.divider()

        st.subheader("⚙️ Menú de Guardar Como (Exportación)")
        formatos_disponibles = {
            "Documento de Word (*.docx)": "docx",
            "Documento habilitado con macros de Word (*.docm)": "docm",
            "Documento de Word 97-2003 (*.doc)": "doc",
            "Plantilla de Word (*.dotx)": "dotx",
            "PDF (*.pdf)": "pdf",
            "Página web (*.html)": "html",
            "Texto sin formato (*.txt)": "txt",
            "Documento XML de Word (*.xml)": "xml",
            "Texto de OpenDocument (*.odt)": "odt"
        }
        
        formato_elegido = st.selectbox("📂 Seleccione el formato de salida deseado:", list(formatos_disponibles.keys()))
        extension_final = formatos_disponibles[formato_elegido]
        
        nombre_puro = os.path.splitext(archivo_subido.name)[0]
        nombre_salida = f"EDITADO_{nombre_puro}.{extension_final}"

        if extension_final in ["docx", "docm", "doc", "dotx"]:
            bytes_archivo = exportar_a_docx(parrafos)
        elif extension_final == "pdf":
            bytes_archivo = exportar_a_pdf(parrafos)
        elif extension_final in ["txt", "odt"]:
            bytes_archivo = exportar_a_txt(parrafos)
        elif extension_final == "xml":
            bytes_archivo = exportar_a_xml(parrafos)
        elif extension_final == "html":
            bytes_archivo = exportar_a_html(parrafos)

        st.download_button(
            label=f"📥 Descargar archivo modificado como .{extension_final}",
            data=bytes_archivo,
            file_name=nombre_salida,
            mime="application/octet-stream"
        )

if __name__ == "__main__":
    main()